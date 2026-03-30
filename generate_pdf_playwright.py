#!/usr/bin/env python3
"""Generate client-ready proposal DOCX files and convert them to PDF with LibreOffice."""

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

BASE_DIR = Path(__file__).parent
OUTPUT_DIR = BASE_DIR / "output"
TODAY_FILE = date.today().strftime("%Y-%m-%d")

HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
UNORDERED_RE = re.compile(r"^\s*[-*+]\s+")
ORDERED_RE = re.compile(r"^\s*\d+\.\s+")
TABLE_ROW_RE = re.compile(r"^\s*\|")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?[\s:-]+\|[\s|:-]*$")
FENCE_RE = re.compile(r"^```")
BLOCKQUOTE_RE = re.compile(r"^\s*>\s?")
HR_RE = re.compile(r"^\s*([-*_]\s*){3,}$")
INLINE_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\((?:[^)]+)\)|<br\s*/?>)", re.IGNORECASE)
METADATA_LINE_RE = re.compile(
    r"^\s*(총\s*대상\s*회사\s*수|대상\s*회사\s*수|number\s*of\s*target\s*companies|target\s*companies|対象会社数|総\s*対象\s*会社\s*数|総\s*대상\s*회사\s*수)\s*:\s*.+$",
    re.IGNORECASE,
)


@dataclass
class FontSpec:
    latin: str
    east_asia: str
    monospace: str = "Courier New"


FONT_BY_LANGUAGE = {
    "en": FontSpec(latin="Aptos", east_asia="Noto Sans"),
    "ja": FontSpec(latin="Noto Sans JP", east_asia="Noto Sans JP"),
    "ko": FontSpec(latin="Noto Sans KR", east_asia="Noto Sans KR"),
    "zh-hant": FontSpec(latin="Noto Sans TC", east_asia="Noto Sans TC"),
    "zh-hans": FontSpec(latin="Noto Sans SC", east_asia="Noto Sans SC"),
}

PROPOSAL_HEADING_DISPLAY = {
    "ja": {
        "Greeting": "ご挨拶",
        "Executive Summary": "要約",
        "What We Found": "現状整理",
        "Why This Matters": "重要な理由",
        "Market And Competitor Snapshot": "市場と競合の概況",
        "Your Hidden Strengths": "貴社の強み",
        "Recommended Direction": "推奨方針",
        "30-60-90 Day Execution Plan": "30-60-90日の実行計画",
        "Recommended Packages": "推奨パッケージ",
        "Pricing Guidance": "料金の考え方",
        "Why Onecation": "Onecationを選ぶ理由",
        "Suggested Next Step": "次のご提案",
        "Closing": "結び",
    },
    "ko": {
        "Greeting": "인사",
        "Executive Summary": "요약",
        "What We Found": "현재 진단",
        "Why This Matters": "왜 중요한가",
        "Market And Competitor Snapshot": "시장 및 경쟁 상황",
        "Your Hidden Strengths": "귀사의 강점",
        "Recommended Direction": "권장 방향",
        "30-60-90 Day Execution Plan": "30-60-90일 실행 계획",
        "Recommended Packages": "추천 패키지",
        "Pricing Guidance": "가격 안내",
        "Why Onecation": "Onecation이 적합한 이유",
        "Suggested Next Step": "다음 단계 제안",
        "Closing": "맺음말",
    },
    "zh-hant": {
        "Greeting": "問候",
        "Executive Summary": "執行摘要",
        "What We Found": "我們的發現",
        "Why This Matters": "為何重要",
        "Market And Competitor Snapshot": "市場與競爭概況",
        "Your Hidden Strengths": "貴公司的隱性優勢",
        "Recommended Direction": "建議方向",
        "30-60-90 Day Execution Plan": "30-60-90 天執行計畫",
        "Recommended Packages": "建議方案",
        "Pricing Guidance": "報價說明",
        "Why Onecation": "選擇 Onecation 的理由",
        "Suggested Next Step": "建議下一步",
        "Closing": "結語",
    },
    "zh-hans": {
        "Greeting": "问候",
        "Executive Summary": "执行摘要",
        "What We Found": "我们的发现",
        "Why This Matters": "为何重要",
        "Market And Competitor Snapshot": "市场与竞争概况",
        "Your Hidden Strengths": "贵公司的隐藏优势",
        "Recommended Direction": "建议方向",
        "30-60-90 Day Execution Plan": "30-60-90 天执行计划",
        "Recommended Packages": "建议方案",
        "Pricing Guidance": "报价说明",
        "Why Onecation": "选择 Onecation 的理由",
        "Suggested Next Step": "建议下一步",
        "Closing": "结语",
    },
}


def normalize_heading_key(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip()).lower()


PROPOSAL_HEADING_DISPLAY_NORMALIZED = {
    language: {normalize_heading_key(key): value for key, value in headings.items()}
    for language, headings in PROPOSAL_HEADING_DISPLAY.items()
}


@dataclass
class MarkdownBlock:
    kind: str
    text: str = ""
    level: int = 0
    ordered: bool = False
    items: list[str] = field(default_factory=list)
    headers: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)


def parse_companies(text: str) -> dict[str, str]:
    lines = text.splitlines(keepends=True)
    result: dict[str, list[str]] = {}
    current: str | None = None

    for line in lines:
        match = re.match(r"^#\s+(.+)", line)
        if match:
            current = match.group(1).strip()
            result[current] = []
            continue
        if current:
            result[current].append(line)

    if not result:
        return {"default": text.strip()}

    return {company: "".join(content).strip() for company, content in result.items()}


def infer_language(text: str) -> str:
    korean = len(re.findall(r"[\uac00-\ud7a3]", text))
    japanese = len(re.findall(r"[\u3040-\u30ff\u31f0-\u31ff]", text))
    han = len(re.findall(r"[\u4e00-\u9fff]", text))
    latin = len(re.findall(r"[A-Za-z]", text))

    if korean >= japanese and korean >= han and korean > 0:
        return "ko"
    if japanese > 0 or han > latin:
        return "ja"
    return "en"


def normalize_language_code(value: str | None) -> str:
    normalized = (value or "").strip().lower()
    if normalized.startswith("english") or normalized in {"en", "en-us", "en-gb"}:
        return "en"
    if normalized.startswith("japanese") or normalized in {"ja", "jp"}:
        return "ja"
    if normalized.startswith("korean") or normalized in {"ko", "kr"}:
        return "ko"
    if normalized.startswith("traditional chinese") or normalized in {"zh-hant", "zh-tw", "tw"}:
        return "zh-hant"
    if normalized.startswith("simplified chinese") or normalized in {"zh-hans", "zh-cn", "cn"}:
        return "zh-hans"
    return normalized or "en"


def normalize_company_name(value: str) -> str:
    text = value.strip().lower()
    text = re.sub(r"\s*\([^)]*\)", "", text)
    text = re.sub(r"\s*（[^）]*）", "", text)
    text = re.sub(r"[^\w\u3040-\u30ff\u3400-\u9fff\uac00-\ud7af]+", "", text)
    return text


def prepare_client_markdown(company_name: str, proposal_md: str) -> str:
    cleaned_lines: list[str] = []
    normalized_company = normalize_company_name(company_name)

    for raw_line in proposal_md.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            if cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")
            continue

        if METADATA_LINE_RE.match(stripped):
            continue

        if stripped.startswith("# ") and normalize_company_name(stripped[2:]) == normalized_company:
            continue

        if stripped.startswith("<!--") and stripped.endswith("-->"):
            continue

        cleaned_lines.append(line)

    while cleaned_lines and cleaned_lines[0] == "":
        cleaned_lines.pop(0)
    while cleaned_lines and cleaned_lines[-1] == "":
        cleaned_lines.pop()

    return "\n".join(cleaned_lines)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_start(lines: list[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and TABLE_ROW_RE.match(lines[index]) is not None
        and TABLE_SEPARATOR_RE.match(lines[index + 1]) is not None
    )


def is_special_block(line: str) -> bool:
    stripped = line.strip()
    return bool(
        not stripped
        or HEADING_RE.match(stripped)
        or FENCE_RE.match(stripped)
        or HR_RE.match(stripped)
        or UNORDERED_RE.match(line)
        or ORDERED_RE.match(line)
        or BLOCKQUOTE_RE.match(line)
        or TABLE_ROW_RE.match(line)
    )


def parse_markdown_blocks(markdown_text: str) -> list[MarkdownBlock]:
    lines = markdown_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: list[MarkdownBlock] = []
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            blocks.append(MarkdownBlock(kind="heading", level=len(heading_match.group(1)), text=heading_match.group(2).strip()))
            index += 1
            continue

        if FENCE_RE.match(stripped):
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not FENCE_RE.match(lines[index].strip()):
                code_lines.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            blocks.append(MarkdownBlock(kind="code", text="\n".join(code_lines).rstrip()))
            continue

        if HR_RE.match(stripped):
            blocks.append(MarkdownBlock(kind="hr"))
            index += 1
            continue

        if is_table_start(lines, index):
            table_lines: list[str] = []
            while index < len(lines) and TABLE_ROW_RE.match(lines[index]):
                table_lines.append(lines[index])
                index += 1

            headers = split_table_row(table_lines[0])
            rows = [
                split_table_row(row_line)
                for row_line in table_lines[2:]
                if row_line.strip() and not TABLE_SEPARATOR_RE.match(row_line)
            ]
            blocks.append(MarkdownBlock(kind="table", headers=headers, rows=rows))
            continue

        if UNORDERED_RE.match(line):
            items: list[str] = []
            while index < len(lines) and UNORDERED_RE.match(lines[index]):
                items.append(UNORDERED_RE.sub("", lines[index]).strip())
                index += 1
            blocks.append(MarkdownBlock(kind="list", ordered=False, items=items))
            continue

        if ORDERED_RE.match(line):
            items: list[str] = []
            while index < len(lines) and ORDERED_RE.match(lines[index]):
                items.append(ORDERED_RE.sub("", lines[index]).strip())
                index += 1
            blocks.append(MarkdownBlock(kind="list", ordered=True, items=items))
            continue

        if BLOCKQUOTE_RE.match(line):
            quote_lines: list[str] = []
            while index < len(lines) and BLOCKQUOTE_RE.match(lines[index]):
                quote_lines.append(BLOCKQUOTE_RE.sub("", lines[index]).strip())
                index += 1
            blocks.append(MarkdownBlock(kind="quote", text=" ".join(part for part in quote_lines if part)))
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines) and not is_special_block(lines[index]):
            paragraph_lines.append(lines[index].strip())
            index += 1
        blocks.append(MarkdownBlock(kind="paragraph", text=" ".join(part for part in paragraph_lines if part)))

    return blocks


def set_font(target: object, font_spec: FontSpec, size_pt: float | None = None, bold: bool | None = None) -> None:
    font = target.font
    font.name = font_spec.latin
    if size_pt is not None:
        font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold

    element = target._element
    rpr = element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_spec.latin)
    rfonts.set(qn("w:hAnsi"), font_spec.latin)
    rfonts.set(qn("w:eastAsia"), font_spec.east_asia)
    rfonts.set(qn("w:cs"), font_spec.east_asia)


def configure_document(document: Document, language: str) -> FontSpec:
    font_spec = FONT_BY_LANGUAGE.get(language, FONT_BY_LANGUAGE["en"])
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    set_font(document.styles["Normal"], font_spec, size_pt=10.5)
    for style_name, size, bold in [
        ("Title", 20, True),
        ("Heading 1", 15, True),
        ("Heading 2", 13, True),
        ("Heading 3", 11.5, True),
        ("List Bullet", 10.5, False),
        ("List Number", 10.5, False),
    ]:
        if style_name in document.styles:
            set_font(document.styles[style_name], font_spec, size_pt=size, bold=bold)

    normal_paragraph = document.styles["Normal"].paragraph_format
    normal_paragraph.line_spacing = 1.4
    normal_paragraph.space_after = Pt(6)
    return font_spec


def display_heading_text(text: str, language: str) -> str:
    normalized = normalize_heading_key(text)
    return PROPOSAL_HEADING_DISPLAY_NORMALIZED.get(language, {}).get(normalized, text)


def add_inline_runs(paragraph, text: str, font_spec: FontSpec) -> None:
    parts = INLINE_TOKEN_RE.split(text)
    for part in parts:
        if not part:
            continue
        if re.fullmatch(r"<br\s*/?>", part, re.IGNORECASE):
            paragraph.add_run().add_break(WD_BREAK.LINE)
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, font_spec, bold=True)
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, FontSpec(font_spec.monospace, font_spec.east_asia))
            continue
        link_match = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", part)
        if link_match:
            label, url = link_match.groups()
            run = paragraph.add_run(f"{label} ({url})")
            set_font(run, font_spec)
            continue
        run = paragraph.add_run(part)
        set_font(run, font_spec)


def add_text_paragraph(document: Document, text: str, font_spec: FontSpec, style: str = "Normal", italic: bool = False) -> None:
    paragraph = document.add_paragraph(style=style)
    add_inline_runs(paragraph, text, font_spec)
    if italic:
        for run in paragraph.runs:
            run.italic = True


def add_table(document: Document, headers: list[str], rows: list[list[str]], font_spec: FontSpec) -> None:
    table = document.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        add_inline_runs(paragraph, header, font_spec)
        for run in paragraph.runs:
            run.bold = True
            set_font(run, font_spec, bold=True)

    for row_index, row in enumerate(rows, start=1):
        padded = (row + [""] * len(headers))[: len(headers)]
        for col_index, cell_text in enumerate(padded):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_inline_runs(paragraph, cell_text.replace("<br>", "\n"), font_spec)
            for run in paragraph.runs:
                set_font(run, font_spec)


def build_docx(company_name: str, proposal_md: str, output_path: Path, *, language_hint: str | None = None) -> Path:
    language = normalize_language_code(language_hint) if language_hint else infer_language(f"{company_name}\n{proposal_md}")
    prepared_markdown = prepare_client_markdown(company_name, proposal_md)
    blocks = parse_markdown_blocks(prepared_markdown)

    document = Document()
    font_spec = configure_document(document, language)

    document.core_properties.title = company_name if company_name != "default" else "Proposal"
    document.core_properties.author = "Onecation"

    if company_name and company_name != "default":
        title = document.add_paragraph(style="Title")
        add_inline_runs(title, company_name, font_spec)

    for block in blocks:
        if block.kind == "heading":
            level = max(1, min(block.level, 3))
            paragraph = document.add_paragraph(style=f"Heading {level}")
            add_inline_runs(paragraph, display_heading_text(block.text, language), font_spec)
            continue

        if block.kind == "paragraph":
            add_text_paragraph(document, block.text, font_spec)
            continue

        if block.kind == "quote":
            add_text_paragraph(document, block.text, font_spec, italic=True)
            continue

        if block.kind == "list":
            style_name = "List Number" if block.ordered else "List Bullet"
            for item in block.items:
                add_text_paragraph(document, item, font_spec, style=style_name)
            continue

        if block.kind == "table":
            add_table(document, block.headers, block.rows, font_spec)
            continue

        if block.kind == "code":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.6)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(block.text)
            set_font(run, FontSpec(font_spec.monospace, font_spec.east_asia), size_pt=9.5)
            continue

        if block.kind == "hr":
            document.add_paragraph("")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def resolve_soffice_binary() -> str | None:
    env_binary = os.environ.get("LIBREOFFICE_BIN")
    if env_binary and Path(env_binary).exists():
        return env_binary

    for candidate in ["soffice", "libreoffice"]:
        resolved = shutil.which(candidate)
        if resolved:
            return resolved
    return None


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> Path | None:
    soffice_binary = resolve_soffice_binary()
    if not soffice_binary:
        print(f"  [WARN] LibreOffice not found. DOCX generated only: {docx_path.name}")
        return None

    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="sales-factory-lo-profile-") as profile_dir, tempfile.TemporaryDirectory(
        prefix="sales-factory-lo-out-"
    ) as out_dir:
        profile_uri = Path(profile_dir).resolve().as_uri()
        command = [
            soffice_binary,
            f"-env:UserInstallation={profile_uri}",
            "--headless",
            "--nologo",
            "--nodefault",
            "--nofirststartwizard",
            "--nolockcheck",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            out_dir,
            str(docx_path),
        ]
        completed = subprocess.run(command, capture_output=True, text=True, check=False)
        generated_pdf = Path(out_dir) / f"{docx_path.stem}.pdf"
        if completed.returncode != 0 or not generated_pdf.exists():
            stderr = (completed.stderr or completed.stdout or "").strip()
            raise RuntimeError(f"LibreOffice conversion failed for {docx_path.name}: {stderr or 'unknown error'}")

        shutil.move(str(generated_pdf), str(pdf_path))
        return pdf_path


def build_output_paths(stem_path: Path) -> tuple[Path, Path]:
    return Path(f"{stem_path}.docx"), Path(f"{stem_path}.pdf")


def resolve_output_stem(out_dir: Path, company_name: str, min_suffix: int = 0) -> Path:
    display_name = company_name if company_name != "default" else "proposal"
    safe_name = re.sub(r'[\\/*?:"<>|]', "_", display_name)
    base_name = f"{safe_name}_proposal_{TODAY_FILE}_playwright"
    candidates = [out_dir / base_name] + [out_dir / f"{base_name}_{index}" for index in range(1, 10)]

    for index, candidate in enumerate(candidates):
        if index < min_suffix:
            continue
        docx_candidate, pdf_candidate = build_output_paths(candidate)
        if not docx_candidate.exists() and not pdf_candidate.exists():
            return candidate

    return candidates[-1]


def generate_company_documents(
    company_name: str,
    proposal_md: str,
    out_dir: Path,
    require_pdf: bool = False,
    language_hint: str | None = None,
) -> tuple[Path, Path | None]:
    stem_path = resolve_output_stem(out_dir, company_name)
    docx_path, pdf_path = build_output_paths(stem_path)

    build_docx(company_name, proposal_md, docx_path, language_hint=language_hint)
    pdf_result = convert_docx_to_pdf(docx_path, pdf_path)
    if require_pdf and pdf_result is None:
        raise RuntimeError("LibreOffice PDF conversion is required but soffice is not available.")
    return docx_path, pdf_result


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate proposal DOCX files and convert them to PDF with LibreOffice.")
    parser.add_argument("--proposal", default=str(BASE_DIR / "proposal.md"))
    parser.add_argument("--company", default=None)
    parser.add_argument("--out", default=str(OUTPUT_DIR))
    parser.add_argument("--language", default=None)
    parser.add_argument("--require-pdf", action="store_true")
    args = parser.parse_args()

    proposal_path = Path(args.proposal)
    if not proposal_path.exists():
        print(f"[ERROR] Proposal file not found: {proposal_path}")
        raise SystemExit(1)

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    proposal_text = proposal_path.read_text(encoding="utf-8")
    companies = parse_companies(proposal_text)
    if args.company:
        companies = {
            company_name: content
            for company_name, content in companies.items()
            if args.company in company_name or company_name in args.company
        }
        if not companies:
            print(f"[ERROR] Company not found: {args.company}")
            raise SystemExit(1)

    print("\n[Sales Factory Document Generator]")
    print(f"  Companies: {len(companies)}")
    print(f"  Output: {out_dir}")
    print(f"  LibreOffice: {resolve_soffice_binary() or 'not found'}\n")

    failed = False
    for company_name, proposal_md in companies.items():
        try:
            docx_path, pdf_path = generate_company_documents(
                company_name,
                proposal_md,
                out_dir,
                require_pdf=args.require_pdf,
                language_hint=args.language,
            )
            print(f"  [OK] {docx_path.name}")
            if pdf_path:
                size_kb = pdf_path.stat().st_size // 1024
                print(f"  [OK] {pdf_path.name} ({size_kb} KB)")
        except PermissionError:
            fallback_stem = resolve_output_stem(out_dir, company_name, min_suffix=1)
            docx_path, pdf_path = build_output_paths(fallback_stem)
            build_docx(company_name, proposal_md, docx_path, language_hint=args.language)
            pdf_result = convert_docx_to_pdf(docx_path, pdf_path)
            print(f"  [OK] {docx_path.name}")
            if pdf_result:
                print(f"  [OK] {pdf_result.name} ({pdf_result.stat().st_size // 1024} KB)")
        except Exception as exc:
            failed = True
            print(f"  [FAILED] {company_name}: {exc}")

    if failed:
        raise SystemExit(1)

    print(f"\nDone: {out_dir}")


if __name__ == "__main__":
    main()
