from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from generate_pdf_playwright import build_docx, build_output_paths, infer_language, parse_companies, parse_markdown_blocks, prepare_client_markdown


class GenerateProposalDocumentsTests(unittest.TestCase):
    def test_parse_companies_splits_h1_blocks(self) -> None:
        text = "# Alpha\n## Executive Summary\nHello\n# Beta\n## Executive Summary\nWorld\n"

        companies = parse_companies(text)

        self.assertEqual(set(companies.keys()), {"Alpha", "Beta"})
        self.assertIn("Hello", companies["Alpha"])
        self.assertIn("World", companies["Beta"])

    def test_prepare_client_markdown_removes_duplicate_heading_and_metadata(self) -> None:
        markdown = "# 栄光情報システム株式会社\n総 대상 회사 수: 1\n\n## 実行概要\n本文입니다.\n"

        prepared = prepare_client_markdown("栄光情報システム株式会社", markdown)

        self.assertNotIn("# 栄光情報システム株式会社", prepared)
        self.assertNotIn("総 대상 회사 수", prepared)
        self.assertIn("## 実行概要", prepared)

    def test_parse_markdown_blocks_detects_headings_lists_and_tables(self) -> None:
        markdown = """## 実行概要
- 첫 번째
- 두 번째

| 항목 | 내용 |
| --- | --- |
| 강점 | SEO |
"""

        blocks = parse_markdown_blocks(markdown)

        self.assertEqual([block.kind for block in blocks], ["heading", "list", "table"])
        self.assertEqual(blocks[0].text, "実行概要")
        self.assertEqual(blocks[1].items, ["첫 번째", "두 번째"])
        self.assertEqual(blocks[2].headers, ["항목", "내용"])
        self.assertEqual(blocks[2].rows[0], ["강점", "SEO"])

    def test_infer_language_detects_primary_script(self) -> None:
        self.assertEqual(infer_language("Marketing proposal for California buyers"), "en")
        self.assertEqual(infer_language("日本市場向けの提案書です"), "ja")
        self.assertEqual(infer_language("한국 시장 제안서입니다"), "ko")

    def test_build_docx_writes_client_ready_document(self) -> None:
        markdown = """## Executive Summary
This is the summary.

## Pricing
| Package | Price |
| --- | --- |
| Core | $5,000 |
"""

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "proposal.docx"
            build_docx("QP Graphics, Inc.", markdown, output_path)

            self.assertTrue(output_path.exists())

            document = Document(output_path)
            texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            self.assertEqual(texts[0], "QP Graphics, Inc.")
            self.assertIn("Executive Summary", texts)
            self.assertIn("This is the summary.", texts)
            self.assertEqual(document.tables[0].cell(1, 0).text, "Core")
            self.assertEqual(document.tables[0].cell(1, 1).text, "$5,000")

    def test_build_docx_localizes_visible_headings_for_target_language(self) -> None:
        markdown = """## Market and Competitor Snapshot
これは要約です。
"""

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "proposal-ja.docx"
            build_docx("栄光情報システム株式会社", markdown, output_path, language_hint="Japanese")

            document = Document(output_path)
            texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            self.assertIn("市場と競合の概況", texts)
            self.assertNotIn("Market and Competitor Snapshot", texts)

    def test_build_output_paths_preserves_company_names_with_periods(self) -> None:
        docx_path, pdf_path = build_output_paths(Path("output/Acme Co., Ltd._proposal_2026-03-30_playwright"))

        self.assertTrue(str(docx_path).endswith("Acme Co., Ltd._proposal_2026-03-30_playwright.docx"))
        self.assertTrue(str(pdf_path).endswith("Acme Co., Ltd._proposal_2026-03-30_playwright.pdf"))


if __name__ == "__main__":
    unittest.main()
